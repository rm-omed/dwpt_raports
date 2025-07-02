import os
from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tempfile import NamedTemporaryFile

from engine.utils import compute_monthly_summary


def fill_template(data: dict, template_filename: str, preview_mode: bool = False) -> str:
    template_path = os.path.abspath(os.path.join("templates", template_filename))
    doc = Document(template_path)

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT
    }

    # 🔁 Replace placeholders in paragraphs
    for p in doc.paragraphs:
        full_text = "".join(run.text for run in p.runs)
        replaced = full_text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in replaced:
                replaced = replaced.replace(placeholder, str(value))

        if replaced != full_text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = replaced
            else:
                p.add_run(replaced)
            if "{{" in replaced:
                print("[WARN] Unreplaced placeholder found:", replaced)

    # 🔁 Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if isinstance(value, str) and placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

        # Structured tables (dynamic rows)
        for key, value in data.items():
            if isinstance(value, list) and value and isinstance(value[0], dict):
                placeholder_row = None
                for row in table.rows:
                    if any("{{" in cell.text for cell in row.cells):
                        placeholder_row = row
                        break
                if not placeholder_row:
                    continue

                tbl = table._tbl
                tbl.remove(placeholder_row._tr)

                for row_data in value:
                    new_row = table.add_row()
                    for j, (col_name, cell_data) in enumerate(row_data.items()):
                        text = str(cell_data.get("text", ""))
                        align = align_map.get(cell_data.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
                        cell = new_row.cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(text)
                        run.font.size = Pt(10)
                        p.alignment = align

                        
        # 📊 Compute summary from weekly data (general case)
    if "week_data" in data and isinstance(data["week_data"], list):
        weeks = data["week_data"]

        def safe_int(x): return int(x) if str(x).isdigit() else 0

        data["monthly_total_visits"] = sum(safe_int(w.get("visits", 0)) for w in weeks)
        data["monthly_total_incidents"] = sum(safe_int(w.get("incidents", 0)) for w in weeks)
        data["monthly_total_repairs"] = sum(safe_int(w.get("repairs", 0)) for w in weeks)

        # ✅ Compute monthly summary (for default12.docx)
    if template_filename == "default12.docx":
        weeks = data.get("weeks", [])
        if isinstance(weeks, list) and len(weeks) == 4:
            summary = compute_monthly_summary(weeks)
            data.update(summary)


    # 📄 Save output file
    if preview_mode:
        temp_file = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        print(f"[DEBUG] Filled template (preview): {template_filename} -> {temp_file.name}")
        return temp_file.name

    number = data.get("number", "").strip() or data.get("num2", "").strip() or "report"
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    filename = f"{number}_{timestamp}.docx"
    output_path = os.path.abspath(os.path.join("data", filename))

    try:
        doc.save(output_path)
    except PermissionError:
        backup = output_path.replace(".docx", f"_backup_{datetime.now().strftime('%H%M%S')}.docx")
        doc.save(backup)
        output_path = backup

    print(f"[DEBUG] Filling template: {template_filename} -> {output_path}")
    return output_path


def generate_preview_text(data: dict, template_filename: str) -> str:
    path = os.path.join("templates", template_filename)
    doc = Document(path)
    output = []

    def replace(text):
        for k, v in data.items():
            placeholder = "{{" + k + "}}"
            if isinstance(v, str):
                text = text.replace(placeholder, v)
        return text

    for para in doc.paragraphs:
        text = replace(para.text)
        if text.strip():
            output.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [replace(cell.text) for cell in row.cells]
            output.append(" | ".join(cells))

    return "\n\n".join(output)
